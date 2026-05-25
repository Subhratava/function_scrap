#!/usr/bin/env python3
"""
Merge placeholder tags like {{T1}}, {{T2}}, ... in a DOCX template using text files,
with custom tag-specific themers for T2 and T3.

Custom themers
--------------
T2 / T2_X:
- The first non-empty line of each T2*.txt file is treated as the heading.
- The heading is exposed through a companion placeholder tag: {{Heading_T2}}, {{Heading_T2_2}}, etc.
- The remaining content is used for {{T2}}, {{T2_2}}, etc.
- Every remaining block from T2*.txt becomes a separate bullet point.

T3:
- Splits T3.txt into comparable blocks.
- For each comparable block, the first 3 non-empty lines are bold:
    1) Sales Comparable X
    2) Address
    3) Date
- The remaining descriptive text stays normal.
"""
import json
from docx.table import Table
from docx.text.paragraph import Paragraph
import argparse
import re
from copy import deepcopy
from pathlib import Path
from typing import Dict, List, Optional
#from Tables.ExcelTable1Copy import read_excel_range, replace_placeholder_with_table
from docx import Document
from docx.oxml import OxmlElement
# from market_approach_azure_map import create_market_approach_map_from_excel
# from image_merge import replace_placeholder_with_image_in_doc

PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_\-]+)\s*\}\}")
T3_HEADER_RE = re.compile(r"^\s*Sales\s+Comparable\s+\d+\s*$", re.IGNORECASE)
T2_TAG_RE = re.compile(r"^T2(?:_\d+)?$", re.IGNORECASE)
HEADING_T2_TAG_RE = re.compile(r"^Heading_T2(?:_\d+)?$", re.IGNORECASE)


class ContentBlock:
    def __init__(self, kind: str, text: str):
        self.kind = kind  # paragraph | bullet
        self.text = text


class TagContent:
    def __init__(self, tag: str, blocks: List[ContentBlock], raw_lines: List[str]):
        self.tag = tag
        self.blocks = blocks
        self.raw_lines = raw_lines


# ---------- Text file parsing ----------
def parse_content_blocks(lines: List[str]) -> List[ContentBlock]:
    blocks: List[ContentBlock] = []
    para_buffer: List[str] = []

    def flush_para():
        nonlocal para_buffer
        if para_buffer:
            text = '\n'.join(para_buffer).strip()
            if text:
                blocks.append(ContentBlock('paragraph', text))
            para_buffer = []

    for line in lines:
        if not line.strip():
            flush_para()
            continue

        bullet_match = re.match(r"^\s*([-*•])\s+(.*)$", line)
        if bullet_match:
            flush_para()
            bullet_text = bullet_match.group(2).strip()
            if bullet_text:
                blocks.append(ContentBlock('bullet', bullet_text))
        else:
            para_buffer.append(line.rstrip())

    flush_para()
    return blocks


def parse_text_file(path: Path) -> TagContent:
    raw = path.read_text(encoding='utf-8-sig')
    lines = raw.splitlines()
    return TagContent(path.stem, parse_content_blocks(lines), lines)


def split_heading_and_body(lines: List[str]):
    heading = ''
    body_start_idx = len(lines)
    for idx, line in enumerate(lines):
        if line.strip():
            heading = line.strip()
            body_start_idx = idx + 1
            break

    body_lines = lines[body_start_idx:]
    while body_lines and not body_lines[0].strip():
        body_lines = body_lines[1:]

    return heading, body_lines


# ---------- DOCX helpers ----------
def copy_paragraph_format(src_paragraph, dst_paragraph):
    src_pPr = src_paragraph._p.pPr
    if src_pPr is not None:
        dst_pPr = deepcopy(src_pPr)
        existing = dst_paragraph._p.pPr
        if existing is not None:
            dst_paragraph._p.remove(existing)
        dst_paragraph._p.insert(0, dst_pPr)


def copy_run_format(src_run, dst_run):
    try:
        dst_run._r.rPr = deepcopy(src_run._r.rPr)
    except Exception:
        pass


def insert_paragraph_after(paragraph, style_name: Optional[str] = None):
    anchor = OxmlElement('w:p')
    paragraph._p.addnext(anchor)
    new_para = paragraph._parent.add_paragraph()
    created_p = new_para._p
    anchor.addnext(created_p)
    anchor.getparent().remove(anchor)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
    paragraph._p = paragraph._element = None


def paragraph_text(paragraph) -> str:
    return ''.join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text


def is_bullet_placeholder_paragraph(text: str) -> Optional[str]:
    # Matches e.g. '{{T2}}', '- {{T2}}', '• {{T2}}'
    m = re.match(r"^\s*(?:[-•*]\s+)?(\{\{\s*([A-Za-z0-9_\-]+)\s*\}\})\s*$", text)
    return m.group(2) if m else None


def clear_paragraph_runs(paragraph):
    for run in paragraph.runs[::-1]:
        paragraph._p.remove(run._r)


def set_paragraph_text_preserve_style(paragraph, new_text: str):
    first_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph_runs(paragraph)
    new_run = paragraph.add_run(new_text)
    if first_run is not None:
        copy_run_format(first_run, new_run)


def add_plain_text(paragraph, text: str):
    paragraph.add_run(text)


# ---------- Custom themers ----------
def is_t2_tag(tag: str) -> bool:
    return bool(T2_TAG_RE.fullmatch(tag))


def is_heading_t2_tag(tag: str) -> bool:
    return bool(HEADING_T2_TAG_RE.fullmatch(tag))


def theme_t2_as_bullets(paragraph, tag_content: TagContent):
    """Each T2/T2_X content block becomes its own paragraph, inheriting template list formatting."""
    current = paragraph
    for block in tag_content.blocks:
        new_p = insert_paragraph_after(current)
        try:
            new_p.style = paragraph.style
        except Exception:
            pass
        copy_paragraph_format(paragraph, new_p)
        lines = block.text.split('\n')
        for idx, line in enumerate(lines):
            if idx:
                new_p.add_run().add_break()
            add_plain_text(new_p, line)
        current = new_p
    remove_paragraph(paragraph)


def split_t3_blocks(raw_lines: List[str]) -> List[List[str]]:
    blocks: List[List[str]] = []
    current: List[str] = []
    for line in raw_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if T3_HEADER_RE.match(stripped):
            if current:
                blocks.append(current)
            current = [stripped]
        else:
            if current:
                current.append(stripped)
            else:
                current = [stripped]
    if current:
        blocks.append(current)
    return blocks


def theme_t3_bold_first_three_lines(paragraph, tag_content: TagContent):
    current = paragraph
    blocks = split_t3_blocks(tag_content.raw_lines)
    for block in blocks:
        new_p = insert_paragraph_after(current)
        try:
            new_p.style = paragraph.style
        except Exception:
            pass
        copy_paragraph_format(paragraph, new_p)

        first_run = paragraph.runs[0] if paragraph.runs else None
        for idx, line in enumerate(block):
            if idx:
                new_p.add_run().add_break()
            run = new_p.add_run(line)
            if first_run is not None:
                copy_run_format(first_run, run)
            if idx < 3:
                run.bold = True

        current = new_p
    remove_paragraph(paragraph)


def build_default_blocks_after(paragraph, tag_content: TagContent):
    current = paragraph
    for block in tag_content.blocks:
        style_name = 'List Bullet' if block.kind == 'bullet' else None
        new_p = insert_paragraph_after(current, style_name=style_name)
        if block.kind != 'bullet':
            try:
                new_p.style = paragraph.style
            except Exception:
                pass
        copy_paragraph_format(paragraph, new_p)
        lines = block.text.split('\n')
        for idx, line in enumerate(lines):
            if idx:
                new_p.add_run().add_break()
            add_plain_text(new_p, line)
        current = new_p
    remove_paragraph(paragraph)


def apply_custom_theme_if_any(paragraph, tag: str, tag_content: TagContent) -> bool:
    tag_upper = tag.upper()
    if is_t2_tag(tag_upper):
        theme_t2_as_bullets(paragraph, tag_content)
        return True
    if tag_upper == 'T3':
        theme_t3_bold_first_three_lines(paragraph, tag_content)
        return True
    return False

def set_cell_text_preserve_format(cell, text):
    if not cell.paragraphs:
        p = cell.add_paragraph()
        p.add_run(str(text))
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = str(text)
        for run in p.runs[1:]:
            run.text = ''
    else:
        p.add_run(str(text))

def update_table_below_header(doc, table_json, header_text="Forecast Summary Canada"):

    # Accept either dict or JSON string
    if isinstance(table_json, str):
        table_json = json.loads(table_json)

    def iter_block_items(parent):
        body = parent.element.body
        for child in body.iterchildren():
            if child.tag.endswith('}p'):
                yield Paragraph(child, parent)
            elif child.tag.endswith('}tbl'):
                yield Table(child, parent)

    target_table = None
    header_found = False

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph) and header_text in block.text.strip():
            header_found = True
            continue

        if header_found and isinstance(block, Table):
            target_table = block
            break

    if target_table is None:
        raise ValueError(f'No table found below header: "{header_text}"')

    # Update header row
    for col_idx, value in enumerate(table_json.get("columns", [])):
        if col_idx < len(target_table.rows[0].cells):
            set_cell_text_preserve_format(target_table.rows[0].cells[col_idx], str(value))

    # Update body rows
    for row_idx, row_data in enumerate(table_json.get("rows", []), start=1):
        if row_idx >= len(target_table.rows):
            break

        row_values = [row_data.get("metric", "")] + row_data.get("values", [])
        for col_idx, value in enumerate(row_values):
            if col_idx < len(target_table.rows[row_idx].cells):
                set_cell_text_preserve_format(target_table.rows[row_idx].cells[col_idx], str(value))

    return target_table


# ---------- Replacement engine ----------
def flatten_blocks_for_inline(tag_content: TagContent) -> str:
    return '\n\n'.join(block.text for block in tag_content.blocks)


def replace_inline_in_paragraph(paragraph, tag_map: Dict[str, TagContent], missing_tags: set):
    text = paragraph_text(paragraph)
    if not text or '{{' not in text:
        return False

    matches = list(PLACEHOLDER_RE.finditer(text))
    if not matches:
        return False

    stripped = text.strip()
    bullet_tag = is_bullet_placeholder_paragraph(stripped)

    # Full placeholder paragraph or bullet placeholder paragraph -> can expand to multiple paragraphs
    if len(matches) == 1 and (stripped == matches[0].group(0) or bullet_tag == matches[0].group(1)):
        tag = matches[0].group(1)
        if tag in tag_map:
            tag_content = tag_map[tag]
            if is_heading_t2_tag(tag):
                replacement = flatten_blocks_for_inline(tag_content)
                set_paragraph_text_preserve_style(paragraph, replacement)
            elif not apply_custom_theme_if_any(paragraph, tag, tag_content):
                if len(tag_content.blocks) == 1:
                    set_paragraph_text_preserve_style(paragraph, tag_content.blocks[0].text)
                else:
                    build_default_blocks_after(paragraph, tag_content)
        else:
            missing_tags.add(tag)
        return True

    # Inline replacement -> flatten content to a single string
    replacements = []
    for m in matches:
        tag = m.group(1)
        if tag in tag_map:
            replacements.append((m.span(), flatten_blocks_for_inline(tag_map[tag])))
        else:
            missing_tags.add(tag)
            replacements.append((m.span(), m.group(0)))

    new_text_parts = []
    pos = 0
    for (start, end), replacement_text in replacements:
        new_text_parts.append(text[pos:start])
        new_text_parts.append(replacement_text)
        pos = end
    new_text_parts.append(text[pos:])

    set_paragraph_text_preserve_style(paragraph, ''.join(new_text_parts))
    return True


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p

    def iter_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    yield from iter_tables(cell.tables)

    yield from iter_tables(doc.tables)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                yield p
            yield from iter_tables(hf.tables)


def build_heading_tag_name(tag: str) -> str:
    return f'Heading_{tag}'


def load_tag_map(content_dir: Path) -> Dict[str, TagContent]:
    tag_map: Dict[str, TagContent] = {}
    for txt_path in sorted(content_dir.glob('*.txt')):
        tag = txt_path.stem
        parsed = parse_text_file(txt_path)

        if is_t2_tag(tag):
            heading, body_lines = split_heading_and_body(parsed.raw_lines)
            body_blocks = parse_content_blocks(body_lines)
            tag_map[tag] = TagContent(tag, body_blocks, body_lines)
            if heading:
                heading_tag = build_heading_tag_name(tag)
                tag_map[heading_tag] = TagContent(
                    heading_tag,
                    [ContentBlock('paragraph', heading)],
                    [heading],
                )
        else:
            tag_map[tag] = parsed
    return tag_map


def normalize_tag_name(raw_tag: str) -> str:
    tag = str(raw_tag).strip().strip("{}").strip("\"'").replace(".txt", "")
    if not re.fullmatch(r"[A-Za-z0-9_\-]+", tag):
        raise ValueError(f"Invalid merge tag: {raw_tag}")
    return tag


def tag_content_from_text(tag: str, text: str) -> TagContent:
    lines = str(text or "").splitlines()
    return TagContent(tag, parse_content_blocks(lines), lines)


def load_tag_map_from_dict(content_by_tag: Dict[str, str]) -> Dict[str, TagContent]:
    tag_map: Dict[str, TagContent] = {}
    for raw_tag, text in content_by_tag.items():
        tag = normalize_tag_name(raw_tag)
        parsed = tag_content_from_text(tag, text)

        if is_t2_tag(tag):
            heading, body_lines = split_heading_and_body(parsed.raw_lines)
            body_blocks = parse_content_blocks(body_lines)
            tag_map[tag] = TagContent(tag, body_blocks, body_lines)
            if heading:
                heading_tag = build_heading_tag_name(tag)
                tag_map[heading_tag] = TagContent(
                    heading_tag,
                    [ContentBlock('paragraph', heading)],
                    [heading],
                )
        else:
            tag_map[tag] = parsed

    return tag_map

# ---------- Table Replacement engine ----------

# def merge_excel_table(doc: Document, excel_path: Path, sheet_name: str, cell_range: str, placeholder: str):
#     """Load an Excel range and replace the matching DOCX placeholder with a formatted table."""
#     data = read_excel_range(Path(excel_path), sheet_name, cell_range)
#     return replace_placeholder_with_table(doc, data, placeholder=placeholder)

# ---------- Delete Placeholders ----------


PLACEHOLDER_PATTERN = re.compile(r"\{\{.*?\}\}")
def delete_paragraph(paragraph):
    """Remove a paragraph from the document safely."""
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def remove_placeholder_lines(doc):
    """
    Remove every paragraph that still contains a placeholder like {{...}}.
    This removes the full line/paragraph, not just the placeholder text.
    """
    # We convert to list first because we may delete items while iterating.
    for p in list(iter_all_paragraphs(doc)):
        try:
            if p.text and PLACEHOLDER_PATTERN.search(p.text):
                delete_paragraph(p)
        except Exception as exc:
            print(f"Warning: could not remove placeholder paragraph: {exc}")


def _context_sort_key(path: Path) -> int:
    """
    Sort Context1.json, Context2.json, Context10.json in numeric order.
    """
    match = re.search(r"Context(\d+)\.json$", path.name, re.IGNORECASE)
    return int(match.group(1)) if match else 10**9


def load_context_jsons_in_order(content_dir: Path) -> list[dict]:
    """
    Load Context1.json, Context2.json, ... from content_dir in numeric order.
    """
    context_files = sorted(
        content_dir.glob("Context*.json"),
        key=_context_sort_key,
    )

    context_payloads = []
    for context_file in context_files:
        with open(context_file, "r", encoding="utf-8") as f:
            context_payloads.append(json.load(f))

    return context_payloads


# def delete_table_below_header(doc, header_text):
#     """
#     Deletes the first table found after the paragraph containing header_text.
#     Returns True if a table was deleted, False otherwise.
#     """

#     def iter_block_items(parent):
#         body = parent.element.body
#         for child in body.iterchildren():
#             if child.tag.endswith('}p'):
#                 yield Paragraph(child, parent)
#             elif child.tag.endswith('}tbl'):
#                 yield Table(child, parent)

#     header_found = False

#     for block in iter_block_items(doc):
#         if isinstance(block, Paragraph) and header_text in block.text.strip():
#             header_found = True
#             continue

#         if header_found and isinstance(block, Table):
#             tbl_element = block._element
#             tbl_element.getparent().remove(tbl_element)
#             return True

#     return False


def delete_table_below_header(doc, header_text):
    """
    Deletes the first table found after the paragraph containing header_text.

    If header_text is NOT of the placeholder format {{text}},
    the header paragraph is also deleted.

    Returns True if a table was deleted, False otherwise.
    """

    def iter_block_items(parent):
        body = parent.element.body
        for child in body.iterchildren():
            if child.tag.endswith('}p'):
                yield Paragraph(child, parent)
            elif child.tag.endswith('}tbl'):
                yield Table(child, parent)

    def is_placeholder_header(text):
        """
        Checks if text is exactly in the format {{something}}.
        Example:
            {{Heading_T2_2}} -> True
            Forecast Summary Canada -> False
        """
        return bool(re.fullmatch(r"\{\{[^{}]+\}\}", text.strip()))

    header_found = False
    header_paragraph = None

    for block in iter_block_items(doc):

        if isinstance(block, Paragraph) and header_text in block.text.strip():
            header_found = True
            header_paragraph = block
            continue

        if header_found and isinstance(block, Table):
            # Delete the table
            tbl_element = block._element
            tbl_element.getparent().remove(tbl_element)

            # Delete header only if it is NOT like {{Heading_T2_2}}
            if not is_placeholder_header(header_text):
                p_element = header_paragraph._element
                p_element.getparent().remove(p_element)

            return True

    return False


def is_valid_address_json(json_data):
    """
    Validates that json_data has the expected structure:

    {
        "addresses": [
            {
                "type": "subject",
                "address": "2727 Miller Cut Off Road, La Porte, TX"
            }
        ]
    }
    """

    if not isinstance(json_data, dict):
        return False

    addresses = json_data.get("addresses")

    if not isinstance(addresses, list):
        return False

    if len(addresses) == 0:
        return False

    for item in addresses:
        if not isinstance(item, dict):
            return False

        if "type" not in item or "address" not in item:
            return False

    return True



def load_json_from_txt(file_path: Path):
    """
    Loads JSON content from a .txt file.
    """
    with file_path.open("r", encoding="utf-8") as f:
        return json.load(f)


# def insert_market_approach_maps(doc, base_dir="."):
#     """
#     Dynamically loads address_*.txt files as JSON and maps them to placeholders:
#         address_summary.txt -> {{map1}}
#         address_dca.txt     -> {{map2}}
#         next file           -> {{map3}}
#         next file           -> {{map4}}
#         ...

#     Then creates market approach map PNGs and replaces placeholders in the doc.
#     """

#     base_path = Path(base_dir)

#     # Preferred ordering for known files
#     preferred_files = [
#         base_path / "address_summary.txt",
#         base_path / "address_dca.txt",
#     ]

#     # Preserve preferred order, then add future files

#     ordered_files = []
#     for index, file_path in enumerate(preferred_files, start=1):
#         if file_path.exists():
#             ordered_files.append((index, file_path))

#     # Process files dynamically
#     for index, json_file in ordered_files:
#         placeholder = f"{{{{map{index}}}}}"

#         json_data = load_json_from_txt(json_file)
        
#         if not is_valid_address_json(json_data):
#             print(f"Skipping {json_file}: invalid address JSON structure")
#             continue


#         output_png = base_path / f"map_{index}.png"
#         output_csv = base_path / f"map_{index}.csv"
#         if not output_png.exists():
#             result = create_market_approach_map_from_excel(
#                 json=json_data,
#                 output_png=str(output_png),
#                 output_csv=str(output_csv),
#             )
#         else:
#             result = {"output_png": str(output_png)}

#         replace_placeholder_with_image_in_doc(
#             doc=doc,
#             image_path=result.get("output_png", str(output_png)),
#             placeholder=placeholder,
#             image_width_inches=7.5,
#         )

def process_document(template_path: Path, content_dir: Path | Dict[str, str], output_path: Path, excel_path: Path | None = None , map: bool = False):
    doc = Document(str(template_path))
    tag_map = (
        load_tag_map_from_dict(content_dir)
        if isinstance(content_dir, dict)
        else load_tag_map(content_dir)
    )
    missing_tags = set()
    context_jsons = [] if isinstance(content_dir, dict) else load_context_jsons_in_order(content_dir)

    forecast_headers = [
        "Forecast Summary Canada",
        "{{Heading_T2_2}}",
    ]
    for context_payload, header_text in zip(context_jsons, forecast_headers):
        update_table_below_header(
            doc,
            context_payload,
            header_text=header_text,
        )
    missing_headers = forecast_headers[len(context_jsons):]
    for header_text in missing_headers:
        delete_table_below_header(doc, header_text)
    
    paragraphs = list(iter_all_paragraphs(doc))
    for p in paragraphs:
        try:
            replace_inline_in_paragraph(p, tag_map, missing_tags)
        except Exception as exc:
            print(f"Warning: could not process a paragraph: {exc}")

    if map:
        pass
        #insert_market_approach_maps(doc)

    #replace tables after paragraphs.
    # if excel_path is not None:
    #     merge_excel_table(
    #     doc,
    #     excel_path=excel_path,
    #     sheet_name="III. Executive Summary",
    #     cell_range="B7:D49",
    #     placeholder="{{Table_1}}",
    # )

    
    # Cleanup any unresolved placeholders and remove their full paragraph/line
    # Delete tables for headers that did not receive context JSON

    remove_placeholder_lines(doc)

    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    print(f"Loaded tag files/placeholders: {', '.join(sorted(tag_map.keys())) if tag_map else '(none)'}")
    if missing_tags:
        print(f"Placeholders found without matching .txt files: {', '.join(sorted(missing_tags))}")
    else:
        print("All placeholders were matched.")


def main():
    parser = argparse.ArgumentParser(description='Replace DOCX placeholders like {{T1}} with text from matching .txt files.')
    parser.add_argument('--template', required=True, help='Path to the template DOCX file')
    parser.add_argument('--content-dir', default='.', help='Directory containing .txt files')
    parser.add_argument('--output', required=True, help='Output DOCX path')
    args = parser.parse_args()

    process_document(Path(args.template), Path(args.content_dir), Path(args.output))


if __name__ == '__main__':
    main()
