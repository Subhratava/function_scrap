#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from io import BytesIO
from datetime import date, datetime, time
from pathlib import Path
from typing import Any, Mapping
from xml.etree import ElementTree
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.styles.colors import COLOR_INDEX
from openpyxl.utils import range_boundaries
from openpyxl.worksheet.worksheet import Worksheet


DEFAULT_TEST_MAP = {
    "summary": ("A1:G8", "table1"),
    "sales": ("A1:E6", "table2"),
}

EXCEL_BORDER_TO_WORD = {
    "hair": "single",
    "thin": "single",
    "medium": "single",
    "thick": "thick",
    "dotted": "dotted",
    "dashed": "dashed",
    "dashDot": "dotDash",
    "dashDotDot": "dotDotDash",
    "double": "double",
    "slantDashDot": "dotDash",
}


DEFAULT_MIN_COLUMN_WIDTH = Inches(0.5)
DEFAULT_MIN_FONT_SIZE_PT = 7.0


def copy_excel_tables_to_docx(
    doc: Document,
    table_map: Mapping[str, tuple[str, str | int]],
    excel_path: str | Path,
    fit_to_page: bool = True,
    max_width_ratio: float = 1.0,
    min_column_width=DEFAULT_MIN_COLUMN_WIDTH,
    min_font_size_pt: float = DEFAULT_MIN_FONT_SIZE_PT,
) -> Document:
    """
    Copy formatted Excel ranges into matching DOCX placeholders.

    table_map format:
        {"Sheet name": ("A1:D5", "table1"), "Other": ("B2:E8", 2)}

    The second tuple item maps to the placeholder tag. Both "table2" and 2
    resolve to {{table2}}. Missing sheets and missing placeholders are skipped.

    When fit_to_page is true, inserted table widths are scaled to fit inside the
    document section's printable page width.
    """
    source_excel = Path(excel_path)

    workbook = load_workbook(source_excel, data_only=True)
    document = doc

    for sheet_name, spec in table_map.items():
        if sheet_name not in workbook.sheetnames:
            continue

        cell_range, table_ref = spec
        placeholder = _placeholder_name(table_ref)
        worksheet = workbook[sheet_name]
        inserted_table = _build_docx_table_from_excel_range(
            document,
            worksheet,
            cell_range,
            fit_to_page=fit_to_page,
            max_width_ratio=max_width_ratio,
            min_column_width=min_column_width,
            min_font_size_pt=min_font_size_pt,
        )
        _replace_placeholder_with_table(document, placeholder, inserted_table)

    return document



def resize_docx_tables_to_fit_page(
    docx_path: str | Path,
    output_path: str | Path | None = None,
    max_width_ratio: float = 1.0,
    min_column_width=DEFAULT_MIN_COLUMN_WIDTH,
    min_font_size_pt: float = DEFAULT_MIN_FONT_SIZE_PT,
) -> Path:
    """
    Resize every table in an existing DOCX so it fits inside page margins.

    This is useful when another module has already created tables and you only
    want a post-processing pass.
    """
    source_docx = Path(docx_path)
    target_docx = Path(output_path) if output_path is not None else source_docx
    document = Document(str(source_docx))
    section = document.sections[0]

    for table in document.tables:
        auto_resize_table_to_page(
            table,
            section,
            max_width_ratio=max_width_ratio,
            min_column_width=min_column_width,
            min_font_size_pt=min_font_size_pt,
        )

    target_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target_docx))
    return target_docx


def auto_resize_table_to_page(
    table,
    section,
    max_width_ratio: float = 1.0,
    min_column_width=DEFAULT_MIN_COLUMN_WIDTH,
    min_font_size_pt: float = DEFAULT_MIN_FONT_SIZE_PT,
) -> None:
    """
    Scale a python-docx table to the printable width of a Word section.

    The function preserves relative column widths, enforces a minimum column
    width, and writes explicit Word table, grid, and cell widths. If the
    minimum-width guard prevents the table from fitting inside the printable
    page width, font sizes greater than ``min_font_size_pt`` are reduced to
    ``min_font_size_pt`` to improve fit/readability without making text
    smaller than the configured floor.
    """
    max_table_width = int(_section_content_width(section) * max_width_ratio)
    if max_table_width <= 0 or not table.rows:
        return

    column_count = len(table.columns)
    if column_count <= 0:
        return

    min_width_emu = int(min_column_width)
    widths = _read_table_column_widths(table)
    if not widths:
        widths = [max_table_width // column_count] * column_count

    # Keep exactly one width per column. This protects existing DOCX tables
    # where the table grid may be incomplete or have extra entries.
    if len(widths) < column_count:
        fallback_width = max(min_width_emu, max_table_width // column_count)
        widths.extend([fallback_width] * (column_count - len(widths)))
    elif len(widths) > column_count:
        widths = widths[:column_count]

    total_width = sum(widths)
    if total_width <= 0:
        return

    should_reduce_font = False
    if total_width > max_table_width:
        should_reduce_font = True
        widths = _scale_widths_to_fit_with_minimum(widths, max_table_width, min_width_emu)

    # If the requested minimum column widths alone exceed the available page
    # width, the table cannot be made to fit by width changes without violating
    # the minimum. Reduce large fonts to the configured floor as a fallback.
    if should_reduce_font or sum(widths) > max_table_width:
        _reduce_table_font_sizes(table, min_font_size_pt=min_font_size_pt)

    _apply_table_widths(table, widths)


def _placeholder_name(table_ref: str | int) -> str:
    if isinstance(table_ref, int):
        return f"table{table_ref}"

    table_ref = str(table_ref).strip()
    if table_ref.startswith("{{") and table_ref.endswith("}}"):
        table_ref = table_ref[2:-2].strip()
    if table_ref.isdigit():
        return f"table{table_ref}"
    return table_ref


def _build_docx_table_from_excel_range(
    document: Document,
    worksheet: Worksheet,
    cell_range: str,
    fit_to_page: bool = True,
    max_width_ratio: float = 1.0,
    min_column_width=DEFAULT_MIN_COLUMN_WIDTH,
    min_font_size_pt: float = DEFAULT_MIN_FONT_SIZE_PT,
):
    min_col, min_row, max_col, max_row = range_boundaries(cell_range)
    row_count = max_row - min_row + 1
    col_count = max_col - min_col + 1

    table = document.add_table(rows=row_count, cols=col_count)
    table.autofit = False
    widths = _excel_column_widths(worksheet, min_col, max_col, min_column_width=min_column_width)
    _apply_table_widths(table, widths)

    for row_offset, excel_row in enumerate(range(min_row, max_row + 1)):
        source_height = worksheet.row_dimensions[excel_row].height
        if source_height:
            table.rows[row_offset].height = Pt(source_height)

        for col_offset, excel_col in enumerate(range(min_col, max_col + 1)):
            excel_cell = worksheet.cell(excel_row, excel_col)
            if isinstance(excel_cell, MergedCell):
                continue

            word_cell = table.cell(row_offset, col_offset)
            _set_cell_text(word_cell, _format_cell_value(excel_cell.value))
            _copy_cell_style(excel_cell, word_cell)

    _apply_merged_cells(table, worksheet, min_col, min_row, max_col, max_row)
    if fit_to_page:
        auto_resize_table_to_page(
            table,
            document.sections[0],
            max_width_ratio=max_width_ratio,
            min_column_width=min_column_width,
            min_font_size_pt=min_font_size_pt,
        )
    copy_excel_images_to_docx_table(
        worksheet,
        table,
        min_col,
        min_row,
        max_col,
        max_row,
    )
    return table


def copy_excel_images_to_docx_table(
    worksheet: Worksheet,
    word_table,
    min_col: int,
    min_row: int,
    max_col: int,
    max_row: int,
) -> None:
    """
    Copy images anchored inside an Excel range into matching DOCX table cells.

    openpyxl exposes worksheet images through a private ``_images`` collection.
    That is currently the practical server-safe path for embedded image access.
    """
    for image in getattr(worksheet, "_images", []):
        position = _image_top_left_position(image)
        if position is None:
            continue

        image_col, image_row = position
        if image_col < min_col or image_col > max_col:
            continue
        if image_row < min_row or image_row > max_row:
            continue

        row_offset = image_row - min_row
        col_offset = image_col - min_col
        try:
            word_cell = word_table.cell(row_offset, col_offset)
        except IndexError:
            continue

        image_stream = _image_to_stream(image)
        if image_stream is None:
            continue

        paragraph = word_cell.paragraphs[0] if word_cell.paragraphs else word_cell.add_paragraph()
        if paragraph.text:
            paragraph = word_cell.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        width, height = _docx_image_size(image, word_cell, word_table.rows[row_offset])
        if width is not None and height is not None:
            run.add_picture(image_stream, width=width, height=height)
        elif width is not None:
            run.add_picture(image_stream, width=width)
        elif height is not None:
            run.add_picture(image_stream, height=height)
        else:
            run.add_picture(image_stream)


def _image_top_left_position(image) -> tuple[int, int] | None:
    anchor = getattr(image, "anchor", None)
    if anchor is None:
        return None

    if isinstance(anchor, str):
        match = re.fullmatch(r"([A-Za-z]+)(\d+)", anchor.strip())
        if not match:
            return None
        col_letters, row_number = match.groups()
        col_number = 0
        for char in col_letters.upper():
            col_number = col_number * 26 + ord(char) - ord("A") + 1
        return col_number, int(row_number)

    marker = getattr(anchor, "_from", None)
    if marker is None:
        return None

    # openpyxl anchor markers are zero-based; worksheet cells are one-based.
    return int(marker.col) + 1, int(marker.row) + 1


def _image_to_stream(image) -> BytesIO | None:
    try:
        data = image._data()
    except Exception:
        return None

    stream = BytesIO(data)
    stream.seek(0)
    return stream


def _docx_image_size(image, word_cell, word_row) -> tuple[Emu | None, Emu | None]:
    cell_width = getattr(word_cell, "width", None)
    row_height = getattr(word_row, "height", None)

    max_width = int(cell_width * 0.92) if cell_width else None
    max_height = int(row_height * 0.88) if row_height else None
    if max_width is not None and max_width <= 0:
        max_width = None
    if max_height is not None and max_height <= 0:
        max_height = None

    image_width_px = getattr(image, "width", None)
    image_height_px = getattr(image, "height", None)
    if not image_width_px or not image_height_px:
        return Emu(max_width) if max_width else None, None

    image_width = int(Inches(float(image_width_px) / 96.0))
    image_height = int(Inches(float(image_height_px) / 96.0))
    if image_width <= 0 or image_height <= 0:
        return None, None

    scale = 1.0
    if max_width is not None:
        scale = min(scale, max_width / image_width)
    if max_height is not None:
        scale = min(scale, max_height / image_height)

    return Emu(max(1, int(image_width * scale))), Emu(max(1, int(image_height * scale)))


def _excel_column_widths(
    worksheet: Worksheet,
    min_col: int,
    max_col: int,
    min_column_width=DEFAULT_MIN_COLUMN_WIDTH,
) -> list[int]:
    widths = []
    min_width_emu = int(min_column_width)
    for excel_col in range(min_col, max_col + 1):
        letter = worksheet.cell(1, excel_col).column_letter
        excel_width = worksheet.column_dimensions[letter].width or 8.43
        widths.append(max(min_width_emu, int(Inches(max(excel_width, 1) * 0.095))))
    return widths


def _scale_widths_to_fit_with_minimum(
    widths: list[int],
    max_table_width: int,
    min_width_emu: int,
) -> list[int]:
    """Scale widths to fit while keeping every column at least min_width_emu."""
    if not widths:
        return widths

    column_count = len(widths)
    if max_table_width <= 0:
        return [max(1, min_width_emu)] * column_count

    # If the minimum itself cannot fit, use the largest equal width that fits.
    # This is the only case where the configured minimum must be relaxed.
    if min_width_emu * column_count > max_table_width:
        equal_width = max(1, max_table_width // column_count)
        return [equal_width] * column_count

    widths = [max(1, int(width)) for width in widths]
    total_width = sum(widths)
    if total_width <= max_table_width:
        return [max(min_width_emu, width) for width in widths]

    scaled = [max(min_width_emu, int(width * max_table_width / total_width)) for width in widths]

    # Reduce columns that are above the minimum until the total fits.
    while sum(scaled) > max_table_width:
        candidates = [idx for idx, width in enumerate(scaled) if width > min_width_emu]
        if not candidates:
            break
        idx = max(candidates, key=lambda i: scaled[i])
        scaled[idx] -= min(scaled[idx] - min_width_emu, sum(scaled) - max_table_width)

    # Add any rounding remainder back to the widest columns.
    remainder = max_table_width - sum(scaled)
    if remainder > 0:
        order = sorted(range(column_count), key=lambda i: widths[i], reverse=True)
        for idx in order:
            if remainder <= 0:
                break
            scaled[idx] += 1
            remainder -= 1

    return scaled


def _reduce_table_font_sizes(table, min_font_size_pt: float = DEFAULT_MIN_FONT_SIZE_PT) -> None:
    """Reduce only explicit run font sizes greater than min_font_size_pt."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        continue
                    current_size_pt = run.font.size.pt
                    if current_size_pt is not None and current_size_pt > min_font_size_pt:
                        run.font.size = Pt(min_font_size_pt)


def _apply_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_preferred_width(table, sum(widths))
    _set_table_grid_widths(table, widths)

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = width
            _set_cell_preferred_width(row.cells[idx], width)


def _section_content_width(section) -> int:
    return int(section.page_width - section.left_margin - section.right_margin)


def _read_table_column_widths(table) -> list[int]:
    widths = []
    grid = table._tbl.tblGrid
    if grid is not None:
        for grid_col in grid.gridCol_lst:
            width = grid_col.get(qn("w:w"))
            if width:
                widths.append(_twips_to_emu(int(width)))

    if widths and sum(widths) > 0:
        return widths

    if table.rows:
        for cell in table.rows[0].cells:
            if cell.width:
                widths.append(int(cell.width))
    return widths


def _set_table_preferred_width(table, width_emu: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(_emu_to_twips(width_emu)))


def _set_table_grid_widths(table, widths: list[int]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)

    for grid_col in list(grid):
        grid.remove(grid_col)

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(_emu_to_twips(width)))
        grid.append(grid_col)


def _set_cell_preferred_width(word_cell, width_emu: int) -> None:
    tc_pr = _get_or_add_tc_pr(word_cell)
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(_emu_to_twips(width_emu)))


def _emu_to_twips(value: int) -> int:
    return round(value / 635)


def _twips_to_emu(value: int) -> int:
    return value * 635


def _set_cell_text(word_cell, text: str) -> None:
    paragraph = word_cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(text)


def _format_cell_value(value: Any) -> str:
    if value is None:
        return ""

    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M")

    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")

    if isinstance(value, time):
        return value.strftime("%H:%M")

    # Numeric values from Excel, e.g. 13.45678
    if isinstance(value, (int, float, Decimal)) and not isinstance(value, bool):
        decimal_value = Decimal(str(value))

        if decimal_value.as_tuple().exponent < -2:
            return str(decimal_value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))

        return str(value)

    # Numeric-looking strings from Excel, e.g. "13.45678"
    if isinstance(value, str):
        stripped_value = value.strip()

        try:
            decimal_value = Decimal(stripped_value)
        except InvalidOperation:
            return value

        # Only round if it has more than 2 decimal places
        if decimal_value.as_tuple().exponent < -2:
            return str(decimal_value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))

        return value

    return str(value)


def _copy_cell_style(excel_cell, word_cell) -> None:
    _copy_fill(excel_cell, word_cell)
    _copy_borders(excel_cell, word_cell)
    _copy_alignment(excel_cell, word_cell)
    _copy_font(excel_cell, word_cell)


def _copy_font(excel_cell, word_cell) -> None:
    font = excel_cell.font
    paragraph = word_cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")

    run.bold = bool(font.bold)
    run.italic = bool(font.italic)
    run.underline = bool(font.underline)
    run.font.strike = bool(font.strike)

    if font.name:
        run.font.name = font.name
    if font.sz:
        run.font.size = Pt(float(font.sz))
    if font.vertAlign == "superscript":
        run.font.superscript = True
    elif font.vertAlign == "subscript":
        run.font.subscript = True

    color = _excel_color_to_hex(font.color, excel_cell.parent.parent)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _copy_alignment(excel_cell, word_cell) -> None:
    alignment = excel_cell.alignment
    paragraph = word_cell.paragraphs[0]

    horizontal_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "centerContinuous": WD_ALIGN_PARAGRAPH.CENTER,
        "distributed": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        "fill": WD_ALIGN_PARAGRAPH.LEFT,
        "general": None,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    paragraph.alignment = horizontal_map.get(alignment.horizontal)

    vertical_map = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "center": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
        "justify": WD_ALIGN_VERTICAL.BOTH,
        "distributed": WD_ALIGN_VERTICAL.BOTH,
    }
    if alignment.vertical in vertical_map:
        word_cell.vertical_alignment = vertical_map[alignment.vertical]

    if alignment.wrap_text:
        _set_word_wrap(word_cell)


def _copy_fill(excel_cell, word_cell) -> None:
    fill = excel_cell.fill
    if fill.fill_type is None:
        return

    color = _excel_color_to_hex(fill.fgColor, excel_cell.parent.parent)
    if not color:
        color = _excel_color_to_hex(fill.bgColor, excel_cell.parent.parent)
    if color:
        _set_cell_shading(word_cell, color)


def _copy_borders(excel_cell, word_cell) -> None:
    borders = {
        "top": excel_cell.border.top,
        "left": excel_cell.border.left,
        "bottom": excel_cell.border.bottom,
        "right": excel_cell.border.right,
    }

    tc_pr = _get_or_add_tc_pr(word_cell)
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)

    tc_borders = OxmlElement("w:tcBorders")
    for side_name, side in borders.items():
        if not side or not side.style:
            continue

        element = OxmlElement(f"w:{side_name}")
        element.set(qn("w:val"), EXCEL_BORDER_TO_WORD.get(side.style, "single"))
        element.set(qn("w:sz"), _border_size(side.style))
        element.set(qn("w:space"), "0")
        element.set(
            qn("w:color"),
            _excel_color_to_hex(side.color, excel_cell.parent.parent) or "000000",
        )
        tc_borders.append(element)

    if len(tc_borders):
        tc_pr.append(tc_borders)


def _border_size(style: str) -> str:
    if style == "thick":
        return "18"
    if style == "medium":
        return "12"
    if style == "hair":
        return "2"
    return "6"


def _set_cell_shading(word_cell, color: str) -> None:
    tc_pr = _get_or_add_tc_pr(word_cell)
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _set_word_wrap(word_cell) -> None:
    tc_pr = _get_or_add_tc_pr(word_cell)
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is not None:
        tc_pr.remove(no_wrap)


def _get_or_add_tc_pr(word_cell):
    tc = word_cell._tc
    tc_pr = tc.get_or_add_tcPr()
    return tc_pr


def _excel_color_to_hex(color, workbook=None) -> str | None:
    if color is None:
        return None
    if color.type == "rgb" and color.rgb:
        rgb = str(color.rgb)
        if len(rgb) == 8:
            return rgb[-6:]
        if len(rgb) == 6:
            return rgb
    if color.type == "theme" and workbook is not None:
        theme_color = _theme_color_to_hex(workbook, color.theme)
        if theme_color:
            return _apply_tint(theme_color, color.tint or 0)
    if color.type == "indexed":
        if color.indexed in (64, 65):
            return None
        try:
            indexed = str(COLOR_INDEX[color.indexed])
        except (IndexError, TypeError):
            return None
        if len(indexed) == 8:
            return indexed[-6:]
        if len(indexed) == 6:
            return indexed
    if color.type == "auto":
        return None
    return None


def _theme_color_to_hex(workbook, theme_index: int) -> str | None:
    theme = getattr(workbook, "loaded_theme", None)
    if not theme:
        return None

    try:
        root = ElementTree.fromstring(theme)
    except ElementTree.ParseError:
        return None

    namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    color_scheme = root.find(".//a:clrScheme", namespace)
    if color_scheme is None:
        return None

    theme_slots = [
        "lt1",
        "dk1",
        "lt2",
        "dk2",
        "accent1",
        "accent2",
        "accent3",
        "accent4",
        "accent5",
        "accent6",
        "hlink",
        "folHlink",
    ]
    try:
        slot_name = theme_slots[int(theme_index)]
    except (IndexError, TypeError, ValueError):
        return None

    slot = color_scheme.find(f"a:{slot_name}", namespace)
    if slot is None:
        return None

    srgb = slot.find(".//a:srgbClr", namespace)
    if srgb is not None and srgb.get("val"):
        return srgb.get("val")

    system = slot.find(".//a:sysClr", namespace)
    if system is not None:
        return system.get("lastClr")

    return None


def _apply_tint(hex_color: str, tint: float) -> str:
    if not tint:
        return hex_color.upper()

    red = int(hex_color[0:2], 16)
    green = int(hex_color[2:4], 16)
    blue = int(hex_color[4:6], 16)

    def tint_channel(channel: int) -> int:
        if tint < 0:
            return round(channel * (1 + tint))
        return round(channel + (255 - channel) * tint)

    return "{:02X}{:02X}{:02X}".format(
        max(0, min(255, tint_channel(red))),
        max(0, min(255, tint_channel(green))),
        max(0, min(255, tint_channel(blue))),
    )


def _apply_merged_cells(
    table,
    worksheet: Worksheet,
    min_col: int,
    min_row: int,
    max_col: int,
    max_row: int,
) -> None:
    for merged_range in worksheet.merged_cells.ranges:
        left, top, right, bottom = range_boundaries(str(merged_range))
        if left < min_col or right > max_col or top < min_row or bottom > max_row:
            continue

        start_cell = table.cell(top - min_row, left - min_col)
        end_cell = table.cell(bottom - min_row, right - min_col)
        start_cell.merge(end_cell)


def _replace_placeholder_with_table(document: Document, placeholder: str, table) -> bool:
    target_pattern = re.compile(r"\{\{\s*" + re.escape(placeholder) + r"\s*\}\}")

    for paragraph in _iter_all_paragraphs(document):
        if target_pattern.search(paragraph.text):
            paragraph._p.addnext(table._tbl)
            _remove_paragraph(paragraph)
            return True

    table._tbl.getparent().remove(table._tbl)
    return False


def _iter_all_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _load_map_from_json(path: str | Path) -> dict[str, tuple[str, str | int]]:
    raw = json.loads(Path(path).read_text(encoding="utf-8"))
    table_map = {}
    for sheet_name, value in raw.items():
        if not isinstance(value, (list, tuple)) or len(value) != 2:
            raise ValueError(
                f"Map value for {sheet_name!r} must be [excel_range, table_ref]."
            )
        table_map[sheet_name] = (str(value[0]), value[1])
    return table_map


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Copy formatted Excel ranges into DOCX {{tableN}} placeholders."
    )
    parser.add_argument("--docx", default="BaseTemplate.docx")
    parser.add_argument("--excel", default="test.xlsx")
    parser.add_argument("--output", default="BaseTemplate_with_tables.docx")
    parser.add_argument(
        "--max-width-ratio",
        type=float,
        default=1.0,
        help="Fraction of printable page width to use for tables. Default: 1.0.",
    )
    parser.add_argument(
        "--min-column-width-inches",
        type=float,
        default=0.35,
        help="Minimum table column width in inches. Default: 0.35.",
    )
    parser.add_argument(
        "--min-font-size",
        type=float,
        default=7.0,
        help="Smallest font size to apply when shrinking oversized tables. Default: 7.",
    )
    parser.add_argument(
        "--no-fit-to-page",
        action="store_true",
        help="Keep Excel-derived widths even if the table exceeds page margins.",
    )
    parser.add_argument(
        "--resize-existing",
        action="store_true",
        help="Resize tables in --docx instead of copying tables from Excel.",
    )
    parser.add_argument(
        "--map-json",
        help=(
            "JSON file like "
            '{"summary": ["A1:G8", "table1"], "sales": ["A1:E6", "table2"]}'
        ),
    )
    args = parser.parse_args()

    if args.resize_existing:
        output = resize_docx_tables_to_fit_page(
            args.docx,
            args.output,
            max_width_ratio=args.max_width_ratio,
            min_column_width=Inches(args.min_column_width_inches),
            min_font_size_pt=args.min_font_size,
        )
        print(f"Saved {output}")
        return

    table_map = _load_map_from_json(args.map_json) if args.map_json else DEFAULT_TEST_MAP
    document = Document(str(args.docx))
    copy_excel_tables_to_docx(
        document,
        table_map,
        args.excel,
        fit_to_page=not args.no_fit_to_page,
        max_width_ratio=args.max_width_ratio,
        min_column_width=Inches(args.min_column_width_inches),
        min_font_size_pt=args.min_font_size,
    )
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    print(f"Saved {output}")


if __name__ == "__main__":
    main()
